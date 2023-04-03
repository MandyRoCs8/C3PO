import datetime
import os
import win32com.client
import openpyxl
import pymsgbox
from docx import Document
from tkcalendar import Calendar
from tkinter import Tk, filedialog, Toplevel, Listbox, Button, END, EXTENDED

# A global variable to store the indices of the selected emails from the email_listbox
selected_email_indices = ()


def input_date(prompt):
    def on_select_date():
        date_dialog.quit()

    date_dialog = Toplevel(root)
    date_dialog.title(prompt)
    date_dialog.lift()
    date_dialog.grab_set()
    calendar = Calendar(date_dialog, selectmode="day",
                        year=datetime.datetime.now().year,
                        month=datetime.datetime.now().month,
                        day=datetime.datetime.now().day)
    calendar.pack(padx=10, pady=10)

    date_select_button = Button(date_dialog, text="Select", command=on_select_date)  # Rename variable here
    date_select_button.pack(pady=5)

    date_dialog.mainloop()
    date_dialog.destroy()

    return calendar.selection_get()


# noinspection PyUnusedLocal
# excel_file and word_doc are not being used below
# It maintains the integrity of the code and prevent any unintended side effects
# The function may be updated in the future to make use of these parameters,
# or they may be needed for specific use cases
def process_email(message, excel_file, word_doc):
    # Load Excel data
    wb = openpyxl.load_workbook(excel_file)
    ws = wb.active

    # Create dictionary with labels and their corresponding cell locations
    labels = {
        'Cherwell': (2, 8),
        'CM Ticket': (3, 8),
        'Anticipated GL Date': (4, 8),
        'Short Summary': (5, 8),
        'Reason': (8, 8),
        'Value': (9, 8),
        'Impact': (10, 8),
        'Additional Information': (11, 8)
    }

    # Create dictionary to store values for each label
    values = {}

    # Extract values from Excel file
    for label, cell in labels.items():
        value = ws.cell(row=cell[0], column=cell[1]).value
        values[label] = value

        # Find the paragraph with the label in the Word document and add the value
        for paragraph in current_doc.paragraphs:

            if label in paragraph.text:
                # Add the value from the Excel file to the end of the paragraph
                paragraph.add_run(" " + str(value))
                print(f"Updated paragraph text: '{paragraph.text}'")

    # Show message box
    pymsgbox.alert("Please choose a location to save the output file in the upcoming dialog.", "Save Output File")

    # Choose output file
    output_file = filedialog.asksaveasfilename(title="Save Output File As", defaultextension=".docx",
                                               filetypes=[("Word files", "*.docx")])
    current_doc.save(output_file)

    if not output_file:
        print("No output file selected.")
        exit()

    # Save the output file
    current_doc.save(output_file)

    # Delete temp Excel file
    os.remove(excel_file)

    print("Data has been extracted and saved to the Word template.")


def process_excel_data(excel_file, doc):
    wb = openpyxl.load_workbook(excel_file)
    ws = wb.active

    # Loop over each row of the Excel sheet, skipping the first row
    for row in ws.iter_rows(min_row=2, values_only=True):
        # Find the paragraph with the label in the Word document and add the value
        for paragraph in doc.paragraphs:
            if row[0] in paragraph.text:
                # Add the value from the Excel file to the end
                paragraph.add_run(" " + str(row[1]))
                print(f"Updated paragraph text: '{paragraph.text}'")

    # Delete temp Excel file
    os.remove(excel_file)
    print("Excel data has been extracted and saved to the Word template.")


def on_select_emails():
    global selected_email_indices
    selected_email_indices = email_listbox.curselection()
    email_selection_dialog.quit()


# Show message
pymsgbox.alert("Please choose a Word template file in the upcoming dialog.", "Choose Word Template File")

# Choose Word template file
root = Tk()
root.withdraw()
template_file = filedialog.askopenfilename(title="Choose Word Template File", filetypes=[("Word files", "*.docx")])
if not template_file:
    print("No Word template file selected.")
    exit()

# Load Word document
current_doc = Document(template_file)


# Connect to Outlook
outlook = win32com.client.Dispatch("Outlook.Application").GetNamespace("MAPI")

# Get the date range from the user
print("Please enter the date range to search for emails with the subject 'C3PO'.")
start_date = input_date("Enter the start date (MM/DD/YYYY): ")
end_date = input_date("Enter the end date (MM/DD/YYYY): ")

# Get emails with the subject 'C3PO'
subject = 'C3PO'
folder = outlook.GetDefaultFolder(6)  # Inbox folder

messages = folder.Items
messages.Sort("[ReceivedTime]", True)  # Sort messages by received time in descending order
restriction = (
    f"[ReceivedTime] >= '{start_date.strftime('%m/%d/%Y %I:%M %p')}' "
    f"AND [ReceivedTime] <= '{end_date.strftime('%m/%d/%Y %I:%M %p')}'"
)

restricted_messages = messages.Restrict(restriction)

emails = [msg for msg in restricted_messages if subject in msg.Subject]

if len(emails) == 0:
    print("No email found with the subject '{}'".format(subject))
else:
    email_selection_dialog = Toplevel(root)
    email_selection_dialog.title("Select Emails")
    email_selection_dialog.geometry("400x300")
    email_selection_dialog.lift()
    email_selection_dialog.grab_set()

    email_listbox = Listbox(email_selection_dialog, width=50, selectmode=EXTENDED)
    email_listbox.pack(pady=10)

    for index, msg in enumerate(emails, start=1):
        received_time = msg.ReceivedTime
        sender_name = msg.SenderName
        email_subject = msg.Subject
        email_listbox.insert(END, f"{index}. {sender_name} - {email_subject} - {received_time}")

    select_button = Button(email_selection_dialog, text="Select", command=on_select_emails)
    select_button.pack(pady=5)

    email_selection_dialog.mainloop()
    email_selection_dialog.destroy()

    # Load Word document
    template_doc = Document(template_file)  # Rename the variable here

    for selected_email_index in selected_email_indices:
        selected_message = emails[selected_email_index]

        # Extract data from email
        attachments = selected_message.Attachments
        selected_excel_file = ''
        for attachment in attachments:
            if attachment.FileName.endswith('.xlsx'):
                attachment.SaveAsFile(os.path.join(os.getcwd(), attachment.FileName))
                selected_excel_file = os.path.join(os.getcwd(), attachment.FileName)
                print("Attachment saved to '{}'".format(selected_excel_file))
                break

        if not selected_excel_file:
            print("No Excel file found in attachments.")
        else:
            # pass selected_message, selected_excel_file, and current_doc to process_email()
            process_email(selected_message, selected_excel_file, current_doc)
